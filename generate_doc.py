# -*- coding: utf-8 -*-
"""生成项目运行流程与用户操作流程 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 * (level + 1))
    return p


def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()  # spacing
    return table


# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('基于 Multi-Agent 架构的\n自动化代码重构与安全审计系统')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('项目运行流程与用户操作手册')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════
# 目录页
# ═══════════════════════════════════════════
add_heading('目 录', level=1)
toc_items = [
    '一、项目概述',
    '二、系统运行流程（技术架构）',
    '    2.1 整体架构图',
    '    2.2 核心组件说明',
    '    2.3 Agent 流水线详解',
    '    2.4 相似度匹配机制',
    '三、用户操作流程',
    '    3.1 环境部署',
    '    3.2 日常使用流程',
    '    3.3 API 调用方式',
    '四、闭环学习机制',
    '五、页面导航地图',
    '六、关键配置说明',
    '七、常见问题',
]
for item in toc_items:
    doc.add_paragraph(item)
doc.add_page_break()

# ═══════════════════════════════════════════
# 一、项目概述
# ═══════════════════════════════════════════
add_heading('一、项目概述', level=1)

add_para('本系统是一个基于 Django 框架的 Web 应用，采用多 Agent 流水线架构，能够自动对源代码（主要为 Python）进行安全漏洞扫描、代码质量分析和历史相似案例匹配。')

add_para('系统通过三个专用 AI Agent 依次处理用户提交的代码，合并去重后生成结构化审计报告，包含问题类型、严重程度、描述、原始代码片段和修复建议。审计结果持久化到数据库，并支持 PDF 导出。')

add_heading('核心能力', level=2)
add_bullet('安全漏洞扫描：基于 bandit 静态分析，覆盖 200+ 漏洞模式')
add_bullet('代码质量分析：LLM 驱动（MiMo/DeepSeek/GPT-4o/Qwen），降级到规则引擎')
add_bullet('历史案例匹配：TF-IDF + 余弦相似度，支持闭环学习')
add_bullet('多入口访问：Web UI、REST API、Django Admin')

doc.add_page_break()

# ═══════════════════════════════════════════
# 二、系统运行流程
# ═══════════════════════════════════════════
add_heading('二、系统运行流程（技术架构）', level=1)

# 2.1 整体架构图
add_heading('2.1 整体架构图', level=2)
add_para('系统运行流程如下图所示：')

flow_text = """
┌─────────────────────────────────────────────────────────────────┐
│                        用户入口                                  │
│  ┌──────────┐    ┌──────────────┐    ┌────────────────┐         │
│  │ Web UI   │    │ REST API     │    │ Django Admin   │         │
│  │ /submit/ │    │ POST /api/   │    │ /admin/        │         │
│  └────┬─────┘    └──────┬───────┘    └────────────────┘         │
│       └────────┬────────┘                                        │
│                ▼                                                  │
│     ┌─────────────────────┐                                      │
│     │   Orchestrator      │  ← 核心调度器                        │
│     └─────────┬───────────┘                                      │
│               ▼                                                  │
│  ┌────────────────────────────────────────────────────────┐      │
│  │                 三 Agent 流水线                          │      │
│  │  ① SecurityAgent  →  ② RefactorAgent  →  ③ Similarity  │      │
│  └────────────────────────────────────────────────────────┘      │
│               ▼                                                  │
│     ┌─────────────────────┐                                      │
│     │  合并 → 去重 → 排序  │  按严重程度排序                      │
│     └─────────┬───────────┘                                      │
│               ▼                                                  │
│     ┌──────────┐  ┌──────────────┐                               │
│     │ 数据库   │  │ 返回响应      │                               │
│     └──────────┘  └──────────────┘                               │
└─────────────────────────────────────────────────────────────────┘
"""
add_code_block(flow_text.strip())

# 2.2 核心组件说明
add_heading('2.2 核心组件说明', level=2)

add_table(
    ['组件', '文件路径', '职责'],
    [
        ['Orchestrator', 'audit/core/orchestrator.py', '调度三个 Agent，合并去重结果'],
        ['SecurityAgent', 'audit/core/agents/security_agent.py', 'bandit 静态安全扫描'],
        ['RefactorAgent', 'audit/core/agents/refactor_agent.py', 'LLM 代码质量分析（降级到规则引擎）'],
        ['SimilarityAgent', 'audit/core/agents/similarity_agent.py', 'TF-IDF 相似案例检索'],
        ['LLMClient', 'audit/core/llm_client.py', '统一 LLM 调用（MiMo/DeepSeek/OpenAI/Qwen）'],
        ['SimilarityEngine', 'audit/core/similarity.py', '相似度计算 + 历史案例库'],
    ]
)

# 2.3 Agent 流水线详解
add_heading('2.3 Agent 流水线详解', level=2)

add_heading('Agent ①：SecurityAgent（安全漏洞扫描）', level=3)
add_bullet('主引擎：bandit 静态分析工具，覆盖 200+ 安全漏洞模式')
add_bullet('检测范围：SQL 注入、硬编码密钥、不安全反序列化、弱加密、命令注入、XSS 等')
add_bullet('降级策略：bandit 不可用时，回退到正则表达式模式匹配')

add_heading('Agent ②：RefactorAgent（代码质量分析）', level=3)
add_bullet('优先策略：调用 LLM（MiMo/DeepSeek/GPT-4o/Qwen）进行智能分析')
add_bullet('降级策略：LLM 不可用时，使用规则引擎（Python AST 分析 + 正则匹配）')
add_bullet('规则检测：函数过长、参数过多、裸 except、魔法数字、TODO 标记、过长行、print 语句等')

add_heading('Agent ③：SimilarityAgent（历史案例匹配）', level=3)
add_bullet('算法：TF-IDF 向量化 + 余弦相似度')
add_bullet('分词：jieba 中文分词进行代码文本预处理')
add_bullet('数据源：数据库中 is_fixed=True 的历史审计记录，或 8 个种子案例')

add_heading('结果合并策略', level=3)
add_bullet('去重规则：按 (issue_type, line_number) 组合去重')
add_bullet('排序规则：critical > high > medium > low > info')

# 2.4 相似度匹配机制
add_heading('2.4 相似度匹配机制', level=2)
add_para('系统使用 TF-IDF（词频-逆文档频率）算法将代码文本向量化，通过余弦相似度计算代码之间的相似程度。jieba 分词器用于对代码中的中文注释和字符串进行分词处理，提升匹配精度。')

add_para('历史案例库来源：', bold=True)
add_bullet('种子案例：通过 python manage.py seed_cases 导入的 8 个常见漏洞模式')
add_bullet('闭环积累：用户修复后标记 is_fixed=True 的审计记录自动进入案例库')

doc.add_page_break()

# ═══════════════════════════════════════════
# 三、用户操作流程
# ═══════════════════════════════════════════
add_heading('三、用户操作流程', level=1)

# 3.1 环境部署
add_heading('3.1 环境部署', level=2)

add_heading('方式一：本地开发环境', level=3)
add_code_block('''# 1. 克隆项目
git clone <repo-url>
cd multi-agent-audit

# 2. 配置环境变量
cp .env.example .env
# 编辑 .env，填入 LLM_API_KEY 等配置

# 3. 安装依赖
pip install -r requirements.txt

# 4. 数据库迁移
python manage.py migrate

# 5. 导入种子案例（可选但推荐）
python manage.py seed_cases

# 6. 启动开发服务器
python manage.py runserver''')

add_heading('方式二：生产环境部署（Ubuntu/Debian）', level=3)
add_code_block('''# 一键部署脚本（安装依赖 + Nginx + Gunicorn + systemd）
sudo bash deploy/deploy.sh''')

add_para('部署脚本自动完成：')
add_bullet('安装系统依赖')
add_bullet('创建 Python 虚拟环境')
add_bullet('执行数据库迁移')
add_bullet('配置 Nginx 反向代理（端口 80 → Gunicorn 8000）')
add_bullet('配置 systemd 服务（自动重启）')

# 3.2 日常使用流程
add_heading('3.2 日常使用流程', level=2)

add_heading('Step 1：访问仪表盘', level=3)
add_para('访问 http://localhost:8000/ 进入系统首页，可查看：')
add_bullet('项目总数、问题总数、修复率统计')
add_bullet('严重程度分布饼图')
add_bullet('问题类型 Top10 柱状图')

add_heading('Step 2：提交代码审计', level=3)
add_para('访问 http://localhost:8000/submit/ 进入代码提交页面：')
add_bullet('填写项目名称')
add_bullet('填写仓库地址（可选）')
add_bullet('选择编程语言（默认 Python）')
add_bullet('粘贴或上传代码内容')
add_bullet('点击"开始审计"按钮')

add_para('提交后系统自动执行三 Agent 流水线，等待数秒至数十秒后展示审计结果：')
add_bullet('每个问题的类型和严重程度')
add_bullet('原始代码片段')
add_bullet('修复建议')
add_bullet('相似历史案例（如有）')

add_heading('Step 3：查看项目列表', level=3)
add_para('访问 http://localhost:8000/projects/ 查看所有已审计项目，点击项目进入详情页。')

add_heading('Step 4：查看审计详情', level=3)
add_para('访问 http://localhost:8000/detail/<id>/ 查看单个文件的审计详情：')
add_bullet('完整的问题列表')
add_bullet('代码高亮显示')
add_bullet('导出 PDF 审计报告')

# 3.3 API 调用方式
add_heading('3.3 API 调用方式', level=2)

add_heading('触发代码审计', level=3)
add_code_block('''POST /api/audit/
Content-Type: application/json

{
    "code": "import os\\nos.system(input())",
    "file_path": "app.py",
    "language": "Python",
    "top_k": 3
}''')

add_heading('搜索相似案例', level=3)
add_code_block('''POST /api/search-similar/
Content-Type: application/json

{
    "code": "SELECT * FROM users WHERE id=" + user_input,
    "top_k": 5
}''')

add_heading('DRF CRUD 接口', level=3)
add_bullet('GET /api/projects/ — 项目列表')
add_bullet('GET /api/code-files/ — 代码文件列表')
add_bullet('GET /api/audit-records/ — 审计记录列表')
add_bullet('支持分页：每页 20 条（可配置）')

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、闭环学习机制
# ═══════════════════════════════════════════
add_heading('四、闭环学习机制', level=1)

add_para('系统支持闭环学习，流程如下：')

add_bullet('审计发现问题 → 保存到 AuditRecord 表（is_fixed=False）')
add_bullet('用户修复代码后 → 标记 is_fixed=True')
add_bullet('已修复记录 → 进入相似案例库（TF-IDF 索引）')
add_bullet('下次审计时 → 自动匹配相似历史案例，提供修复参考')

add_para('这意味着系统使用越多，案例库越丰富，相似匹配越精准。')

doc.add_page_break()

# ═══════════════════════════════════════════
# 五、页面导航地图
# ═══════════════════════════════════════════
add_heading('五、页面导航地图', level=1)

add_table(
    ['URL 路径', '功能', '说明'],
    [
        ['/', '仪表盘', '统计概览 + 图表展示'],
        ['/submit/', '提交代码', '核心功能入口，提交代码进行审计'],
        ['/projects/', '项目列表', '查看所有已审计项目'],
        ['/detail/<id>/', '审计详情', '单文件审计结果详情'],
        ['/export-pdf/<id>/', '导出 PDF', '下载 PDF 格式审计报告'],
        ['/admin/', '后台管理', 'Django Admin 管理界面'],
        ['/api/audit/', '审计 API', '程序化触发代码审计'],
        ['/api/search-similar/', '相似搜索 API', '查找历史相似案例'],
        ['/api/projects/', '项目 CRUD', '项目增删改查接口'],
        ['/api/code-files/', '文件 CRUD', '代码文件增删改查接口'],
        ['/api/audit-records/', '记录 CRUD', '审计记录增删改查接口'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 六、关键配置说明
# ═══════════════════════════════════════════
add_heading('六、关键配置说明', level=1)

add_heading('环境变量（.env 文件）', level=2)

add_table(
    ['变量名', '说明', '示例值'],
    [
        ['DJANGO_SECRET_KEY', 'Django 密钥（生产环境必填）', 'your-secret-key-here'],
        ['DJANGO_DEBUG', '调试模式开关', 'True / False'],
        ['DJANGO_ALLOWED_HOSTS', '允许的主机名（逗号分隔）', 'localhost,127.0.0.1'],
        ['LLM_API_KEY', 'LLM 服务 API 密钥', 'sk-xxx'],
        ['DJANGO_SECURE_SSL', '生产环境 SSL 重定向', 'True / False'],
    ]
)

add_heading('LLM 配置（config/settings.py）', level=2)

add_table(
    ['配置项', '说明', '默认值'],
    [
        ['LLM_CONFIG.provider', 'LLM 服务商', 'mimo'],
        ['LLM_CONFIG.api_key', 'API 密钥（从环境变量读取）', '—'],
        ['LLM_CONFIG.base_url', 'API 端点地址', '自动按 provider 设置'],
        ['LLM_CONFIG.model', '模型名称', 'mimo-v2.5-pro'],
    ]
)

add_para('支持的 LLM 服务商：', bold=True)
add_bullet('mimo — 小米 MiMo（Anthropic 兼容接口）')
add_bullet('deepseek — DeepSeek（OpenAI 兼容接口）')
add_bullet('openai — GPT-4o（OpenAI API）')
add_bullet('qwen — 阿里通义千问（DashScope 兼容模式）')

doc.add_page_break()

# ═══════════════════════════════════════════
# 七、常见问题
# ═══════════════════════════════════════════
add_heading('七、常见问题', level=1)

qa_list = [
    ('Q: LLM 不可用时系统还能工作吗？',
     'A: 可以。RefactorAgent 会自动降级到规则引擎（AST + 正则匹配），SecurityAgent 使用 bandit 独立工作，SimilarityAgent 使用 TF-IDF 算法，均不依赖 LLM。'),
    ('Q: 如何切换 LLM 服务商？',
     'A: 修改 config/settings.py 中的 LLM_CONFIG.provider 字段，并确保对应的 API Key 已配置在环境变量中。'),
    ('Q: 种子案例有什么用？',
     'A: 种子案例（8 个常见漏洞模式）作为相似匹配的初始数据源。运行 python manage.py seed_cases 导入。随着使用积累，系统会用真实审计记录替代种子案例。'),
    ('Q: 审计结果如何导出？',
     'A: 在审计详情页点击"导出 PDF"按钮，系统通过 xhtml2pdf 生成可下载的 PDF 报告。'),
    ('Q: 支持哪些编程语言？',
     'A: 目前主要支持 Python。SecurityAgent 的 bandit 工具专注于 Python，RefactorAgent 的规则引擎也是 Python AST 分析。LLM 分析理论上支持多语言，但效果最佳为 Python。'),
]

for q, a in qa_list:
    add_para(q, bold=True)
    add_para(a)
    doc.add_paragraph()  # spacing

# ── 保存 ──
output_path = 'D:/multi-agent-audit/项目运行流程与用户操作手册.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
